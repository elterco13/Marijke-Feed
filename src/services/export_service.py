"""
Export service — handles CSV, XLSX, DOCX, and PDF exports.
"""

from __future__ import annotations

import io
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)

from src.db.models import Result

logger = logging.getLogger(__name__)


def results_to_dataframe(results: list[Result]) -> pd.DataFrame:
    """Convert Result ORM objects to a clean DataFrame."""
    rows = []
    for r in results:
        authors = r.get_authors()
        taxa = r.get_taxa()
        species = r.get_species()
        rows.append({
            "Title": r.title,
            "Category": r.category or "",
            "Subcategory": r.subcategory or "",
            "Species": "; ".join(species),
            "Taxa": "; ".join(taxa),
            "Journal/Outlet": r.journal_or_outlet or "",
            "Source": r.source_name or r.source_connector,
            "Published": r.published_at.strftime("%Y-%m-%d") if r.published_at else "",
            "DOI": r.doi or "",
            "URL": r.url or "",
            "Authors": "; ".join(authors),
            "Abstract": (r.abstract_or_summary or "")[:500],
            "Relevance Score": r.relevance_score,
            "Content Type": r.content_type or "",
            "Is Preprint": "Yes" if r.is_preprint else "No",
            "Is Saved": "Yes" if r.is_saved else "No",
            "Is Irrelevant": "Yes" if r.is_irrelevant else "No",
        })
    return pd.DataFrame(rows)


def export_csv(results: list[Result]) -> bytes:
    """Export results as CSV bytes."""
    df = results_to_dataframe(results)
    buf = io.BytesIO()
    df.to_csv(buf, index=False, encoding="utf-8-sig")
    return buf.getvalue()


def export_xlsx(results: list[Result]) -> bytes:
    """Export results as XLSX bytes with formatting."""
    df = results_to_dataframe(results)
    buf = io.BytesIO()

    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Results")

        ws = writer.sheets["Results"]

        # Auto-fit columns (approximate)
        for col_cells in ws.columns:
            max_length = max((len(str(cell.value or "")) for cell in col_cells), default=10)
            col_letter = col_cells[0].column_letter
            ws.column_dimensions[col_letter].width = min(max_length + 4, 60)

        # Style header row
        from openpyxl.styles import Font, PatternFill, Alignment
        header_fill = PatternFill(start_color="0e7490", end_color="0e7490", fill_type="solid")
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

    return buf.getvalue()


def export_docx(results: list[Result], profile_name: str = "Search Profile") -> bytes:
    """Export results as a formatted DOCX document."""
    doc = Document()

    # Title
    title = doc.add_heading("🐠 Aquarium Science Monitor", 0)
    title.runs[0].font.color.rgb = RGBColor(0x0e, 0x74, 0x90)

    # Subtitle
    subtitle = doc.add_paragraph(f"Profile: {profile_name} | Generated: {datetime.utcnow().strftime('%Y-%m-%d')}")
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph("")

    for i, r in enumerate(results, 1):
        # Item heading
        h = doc.add_heading(f"{i}. {r.title}", level=2)
        h.runs[0].font.size = Pt(13)

        # Metadata paragraph
        meta_parts = []
        if r.journal_or_outlet:
            meta_parts.append(f"Journal: {r.journal_or_outlet}")
        if r.published_at:
            meta_parts.append(f"Published: {r.published_at.strftime('%Y-%m-%d')}")
        if r.doi:
            meta_parts.append(f"DOI: {r.doi}")
        meta_parts.append(f"Score: {r.relevance_score:.1f}")
        if r.category:
            meta_parts.append(f"Category: {r.category}")
        if r.is_preprint:
            meta_parts.append("[PREPRINT]")

        meta_p = doc.add_paragraph(" | ".join(meta_parts))
        meta_p.runs[0].font.size = Pt(9)
        meta_p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        # Authors
        authors = r.get_authors()
        if authors:
            auth_p = doc.add_paragraph("Authors: " + "; ".join(authors[:5]))
            auth_p.runs[0].font.size = Pt(9)
            auth_p.runs[0].italic = True

        # Species/taxa
        taxa = r.get_taxa()
        if taxa:
            taxa_p = doc.add_paragraph("Taxa: " + ", ".join(taxa[:8]))
            taxa_p.runs[0].font.size = Pt(9)

        # Abstract
        if r.abstract_or_summary:
            abs_p = doc.add_paragraph(r.abstract_or_summary[:600])
            abs_p.runs[0].font.size = Pt(10)

        # Link
        if r.url:
            link_p = doc.add_paragraph(f"Link: {r.url}")
            link_p.runs[0].font.size = Pt(9)
            link_p.runs[0].font.color.rgb = RGBColor(0x0e, 0x74, 0x90)

        doc.add_paragraph("")  # spacer

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(results: list[Result], profile_name: str = "Search Profile") -> bytes:
    """Export results as PDF using ReportLab."""
    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Aquarium Science Monitor",
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=20,
        textColor=colors.HexColor("#0e7490"),
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#64748b"),
        spaceAfter=20,
    )
    item_title_style = ParagraphStyle(
        "ItemTitle",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=4,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#64748b"),
        spaceAfter=2,
    )
    abstract_style = ParagraphStyle(
        "Abstract",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.HexColor("#374151"),
        spaceAfter=4,
        leading=14,
    )

    story = []

    # Header
    story.append(Paragraph("🐠 Aquarium Science Monitor", title_style))
    story.append(Paragraph(
        f"Profile: {profile_name} | Generated: {datetime.utcnow().strftime('%Y-%m-%d')} | {len(results)} results",
        subtitle_style,
    ))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#0e7490")))
    story.append(Spacer(1, 0.3 * cm))

    for i, r in enumerate(results, 1):
        # Title
        safe_title = _safe(f"{i}. {r.title}")
        story.append(Paragraph(safe_title, item_title_style))

        # Meta line
        meta_parts = []
        if r.journal_or_outlet:
            meta_parts.append(r.journal_or_outlet)
        if r.published_at:
            meta_parts.append(r.published_at.strftime("%Y-%m-%d"))
        if r.doi:
            meta_parts.append(f"DOI: {r.doi}")
        meta_parts.append(f"Score: {r.relevance_score:.1f}")
        if r.category:
            meta_parts.append(r.category)
        if r.is_preprint:
            meta_parts.append("[PREPRINT]")
        story.append(Paragraph(" · ".join(meta_parts), meta_style))

        # Abstract
        if r.abstract_or_summary:
            abstract = r.abstract_or_summary[:500]
            if len(r.abstract_or_summary) > 500:
                abstract += "..."
            story.append(Paragraph(_safe(abstract), abstract_style))

        # URL
        if r.url:
            story.append(Paragraph(f'<link href="{r.url}" color="#0e7490">{r.url[:80]}</link>', meta_style))

        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e2e8f0")))
        story.append(Spacer(1, 0.2 * cm))

    doc.build(story)
    return buf.getvalue()


def _safe(text: str) -> str:
    """Escape XML special characters for ReportLab."""
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
    )
